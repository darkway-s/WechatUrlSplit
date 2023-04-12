from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
def convert_to_docx(text_nodes, format_levels, outputfile):
    doc = Document()
    current_level = None
    for node in text_nodes:
        paragraph = doc.add_paragraph()

        # Step 6: Loop through each format slot and apply a font size to it based on its level.
        slot_name = node.parent.name
        # if format_levels[slot_name] exists
        if slot_name in format_levels:
            level = format_levels[slot_name]
            font_size = 12 + (level * 8)
        else:
            level = 0
            font_size = 6

        # Step 7: Check if the current level is different from the previous level.
        if current_level is not None and current_level < level:
            doc.add_paragraph("----我是一条分割线----")

        run = paragraph.add_run(str(node))
        run.font.size = Pt(font_size)
        
        # Step 8: Update the current level.
        current_level = level
        
    doc.save(outputfile)

def StyleSlot(html_content, outputfile, L=5):
    """
    Convert an HTML file to a Word document, using different font sizes for different HTML tags.
    :param html_content: The HTML content to convert.
    :param outputfile: The path to the output Word document.
    :param L: The number of different font sizes to use.
    :return: A string indicating the status of the conversion.
    """
    # Step 1: Parse the HTML content and extract all the text nodes.
    
    soup = BeautifulSoup(html_content, 'html.parser')
    text_nodes = soup.find_all(string=True)

    # Step 2: Group the text nodes into different format slots based on the formatting applied to them.
    format_slots = {}
    for node in text_nodes:
        if node.parent.name in format_slots:
            format_slots[node.parent.name].append(node)
        else:
            format_slots[node.parent.name] = [node]

    # Step 3: Calculate the length of each text node in each format slot, and sum them up to get the total length of each format slot.
    slot_lengths = {}
    for slot_name, nodes in format_slots.items():
        slot_lengths[slot_name] = sum(len(str(node)) for node in nodes)

    # Step 4: Assign a level to each format slot based on the total length of the text in it.
    sorted_slots = sorted(slot_lengths.items(), key=lambda x: x[1], reverse=True)

    # Step 5: Only keep the L most common format slots.
    sorted_slots = sorted_slots[:L]

    format_levels = {slot[0]: i+1 for i, slot in enumerate(sorted_slots)}

    # Step 6: Use python-docx to convert the HTML to a Word document.
    convert_to_docx(text_nodes, format_levels, outputfile)
    return 'Done!'


if '__main__' == __name__:
    inputfile = 'test.html'
    with open(inputfile, 'r', encoding='utf-8') as f:
        html_content = f.read()
    outputfile = 'output.docx'
    st = StyleSlot(html_content, outputfile)
    print(st)